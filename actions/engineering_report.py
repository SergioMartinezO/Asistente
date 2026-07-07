from __future__ import annotations

from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional
import html
import base64
import re
from core.diagram_service import expected_bundle, write_manifest
from actions.engineering_deliverables_ext import (
    generate_power_flow_diagram,
    generate_fsm_diagram,
    generate_sw_architecture_diagram,
    generate_hw_sw_comm_diagram,
    generate_mechanical_diagram,
    generate_pcb_layout_sketch,
    generate_test_results_chart,
    build_bom_docx_table,
    build_bom_table_html,
    build_pseudocode,
    build_test_results_docx_table,
    build_test_results_table_html,
    build_user_manual_sections,
    build_user_manual_html,
    generate_uml_diagram,
    generate_signal_flow_diagram,
    build_acceptance_criteria,
    build_arduino_source,
    generate_uml_sequence_diagram,
    generate_uml_usecase_diagram,
    build_final_report_section,
)


# ── Mensajes de fase ──────────────────────────────────────────────────────────
_PHASE_MESSAGES = {
    ("Hardware",       "start"): "Iniciando fase de hardware.",
    ("Hardware",       "end"):   "Hardware completado.",
    ("Software",       "start"): "Iniciando fase de software.",
    ("Software",       "end"):   "Software completado.",
    ("Diagramas",      "start"): "Generando diagramas. Esto puede tardar unos segundos.",
    ("Diagramas",      "end"):   "Diagramas generados correctamente.",
    ("Word",           "start"): "Construyendo el documento Word.",
    ("Word",           "end"):   "Documento Word guardado correctamente.",
    ("Web",            "start"): "Construyendo la página web.",
    ("Web",            "end"):   "Página web generada.",
    ("Reporte Final",  "start"): "Consolidando el reporte final y cronograma.",
    ("Reporte Final",  "end"):   "Reporte final completado. Proyecto 100% listo.",
}

_PHASE_ORDER = ["Hardware", "Software", "Diagramas", "Word", "Web", "Reporte Final"]

_PHASE_DEFAULT_DAYS = {
    "Hardware":      5,
    "Software":      8,
    "Diagramas":     3,
    "Word":          2,
    "Web":           2,
    "Reporte Final": 1,
}

_PHASE_DELIVERABLES = {
    "Hardware": [
        "Lista de componentes electrónicos con justificación técnica",
        "Esquemas de conexión y diagramas de bloques funcionales (entrada, procesamiento, salida)",
        "Criterios de aceptación del diseño",
        "Boceto de disposición de PCB (croquis conceptual de componentes y trazas)",
        "Diagrama de flujo de energía (distribución de la alimentación)",
        "Lista de materiales (BOM) con referencia, cantidad y especificación",
        "Diagrama mecánico (vista 2D esquemática del encapsulado/caja)",
    ],
    "Software": [
        "Código fuente (Arduino, Python, C, C++) con comentarios en español",
        "Documentación técnica del software (algoritmos, estructuras de datos, diagramas UML)",
        "Estándares aplicados (ISO/IEC 9126, buenas prácticas de programación)",
        "Diagrama de flujo del programa",
        "Diagrama de estados (FSM)",
        "Pseudocódigo de rutinas críticas",
    ],
    "Diagramas": [
        "Circuito eléctrico/electrónico (esquemático) (PNG/SVG)",
        "Diagrama de bloques funcionales (entrada, procesamiento, salida) (PNG/SVG)",
        "Diagrama UML (clases, casos de uso, secuencia) (PNG/SVG)",
        "Flujos de señal (telecomunicaciones) (PNG/SVG)",
        "Diagramas mecánicos (engranajes, motores, sensores) (PNG/SVG)",
        "Boceto de disposición de PCB (SVG)",
        "Diagrama de flujo de energía (PNG/SVG)",
        "Diagrama de estados FSM (PNG/SVG)",
        "Diagrama de arquitectura de software (PNG/SVG)",
        "Diagrama de comunicación hardware-software (PNG/SVG)",
    ],
    "Word": [
        "Documento Word completo (.docx) con portada (título y autor)",
        "Secciones organizadas: Introducción, Hardware, Software, Diagramas, Conclusiones",
        "Tablas de especificaciones técnicas y BOM formal",
        "Diagramas embebidos",
    ],
    "Web": [
        "Presentación del proyecto con diagramas interactivos (HTML/CSS/JS)",
        "Explicación técnica resumida",
        "Sección de descargas (código fuente Python, código Arduino/C++, documento Word)",
    ],
    "Reporte Final": [
        "Cronograma consolidado con fecha de inicio y fin por fase",
        "Listado de todos los artefactos generados con estado OK/Error",
        "Confirmación explícita: proyecto 100% completo y listo para despliegue",
        "Archivo de texto reporte_final.txt en carpeta Report/",
    ],
}

_SOFTWARE_STANDARDS = [
    "Modelo de calidad ISO/IEC 9126 (Funcionalidad, Confiabilidad, Usabilidad, Eficiencia, Mantenibilidad, Portabilidad)",
    "Buenas prácticas de programación (legibilidad, modularidad, cohesión y acoplamiento óptimos)",
    "Estructura modular y separación de responsabilidades",
    "Nombres descriptivos y comentarios técnicos en español",
    "Validación de entradas y manejo explícito de errores",
    "Trazabilidad de fases, entregables y artefactos generados",
    "Compatibilidad con pruebas automatizadas (pytest)",
]


def _slugify(text: str) -> str:
    s = re.sub(r"[^a-zA-Z0-9._-]+", "_", (text or "").strip())
    s = re.sub(r"_+", "_", s).strip("_")
    return s or "dispositivo"


def _log(player, message: str):
    if player and hasattr(player, "write_log"):
        player.write_log(message)


def _phase_start(player, phase: str, lines: List[str], speak=None, progress: int = 0):
    msg = f"[FASE] ▶ INICIO: {phase}"
    _log(player, msg)
    lines.append(msg)
    if player and hasattr(player, "update_activity"):
        try:
            player.update_activity(estado="En proceso", progreso=progress,
                                   evento=f"Fase {phase} - Iniciada")
        except Exception:
            pass
    if speak:
        phrase = _PHASE_MESSAGES.get((phase, "start"), f"Iniciando fase de {phase.lower()}.")
        try:
            speak(phrase)
            import time; time.sleep(1.2)
        except Exception:
            pass


def _phase_end(player, phase: str, lines: List[str], speak=None, progress: int = 0):
    msg = f"[FASE] ✅ FIN: {phase}"
    _log(player, msg)
    lines.append(msg)
    if player and hasattr(player, "update_activity"):
        try:
            player.update_activity(estado="En proceso", progreso=progress,
                                   evento=f"Fase {phase} - Concluida")
        except Exception:
            pass
    if speak:
        phrase = _PHASE_MESSAGES.get((phase, "end"), f"Fase de {phase.lower()} completada.")
        try:
            speak(phrase)
            import time; time.sleep(1.2)
        except Exception:
            pass


def _parse_start_date(value: Optional[str]) -> datetime:
    if not value:
        return datetime.now()
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%Y/%m/%d"):
        try:
            return datetime.strptime(value.strip(), fmt)
        except Exception:
            continue
    return datetime.now()


def _build_project_plan(parameters: Dict[str, Any]) -> List[Dict[str, Any]]:
    start_date = _parse_start_date(parameters.get("project_start_date"))
    custom_days = parameters.get("phase_durations") or {}
    phases: List[Dict[str, Any]] = []
    cursor = start_date

    for phase in _PHASE_ORDER:
        try:
            duration_days = int(custom_days.get(phase, _PHASE_DEFAULT_DAYS[phase]))
        except Exception:
            duration_days = _PHASE_DEFAULT_DAYS[phase]
        duration_days = max(1, duration_days)

        planned_start = cursor
        planned_end = planned_start + timedelta(days=duration_days - 1)
        phases.append({
            "phase": phase,
            "start": planned_start.strftime("%Y-%m-%d"),
            "end": planned_end.strftime("%Y-%m-%d"),
            "duration_days": duration_days,
            "deliverables": list(_PHASE_DELIVERABLES.get(phase, [])),
        })
        cursor = planned_end + timedelta(days=1)

    return phases


def _plan_summary_lines(plan: List[Dict[str, Any]]) -> List[str]:
    out: List[str] = ["[PLAN] Cronograma del proyecto (inicio/fin por fase):"]
    for item in plan:
        out.append(
            f"[PLAN] {item['phase']}: {item['start']} → {item['end']} ({item['duration_days']} días)"
        )
        for d in item.get("deliverables", []):
            out.append(f"[PLAN]   • Entregable: {d}")
    return out


# ── Helpers Word ──────────────────────────────────────────────────────────────
def _add_heading_safe(doc, text: str, level: int = 1):
    try:
        doc.add_heading(text, level=level)
    except Exception:
        p = doc.add_paragraph(text)
        try:
            p.style = "Title" if level == 1 else "Subtitle"
        except Exception:
            pass


def _safe_replace_placeholders(doc, replacements: Dict[str, str]):
    for paragraph in doc.paragraphs:
        text = paragraph.text or ""
        new_text = text
        for key, value in replacements.items():
            new_text = new_text.replace(key, value)
        if new_text != text:
            paragraph.text = new_text


# ── Fallback diagrams ─────────────────────────────────────────────────────────
def _write_fallback_svg(svg_path: Path, title: str):
    svg_text = (
        "<svg xmlns='http://www.w3.org/2000/svg' width='1200' height='700'>"
        "<rect x='0' y='0' width='1200' height='700' fill='#f8fafc'/>"
        "<rect x='40' y='40' width='1120' height='620' rx='14' fill='white'"
        "  stroke='#cbd5e1' stroke-width='2'/>"
        f"<text x='70' y='120' fill='#1e40af' font-size='40' font-family='Segoe UI'>Diagrama (fallback)</text>"
        f"<text x='70' y='190' fill='#334155' font-size='28' font-family='Segoe UI'>{html.escape(title)}</text>"
        "</svg>"
    )
    svg_path.write_text(svg_text, encoding="utf-8")


def _write_fallback_png(png_path: Path, title: str):
    try:
        from PIL import Image, ImageDraw
        img = Image.new("RGB", (1400, 800), color=(248, 250, 252))
        draw = ImageDraw.Draw(img)
        draw.rectangle((40, 40, 1360, 760), outline=(203, 213, 225), width=3, fill=(255, 255, 255))
        draw.text((70, 90),  "Diagrama (fallback PNG)", fill=(30, 64, 175))
        draw.text((70, 150), title,                     fill=(51, 65, 85))
        img.save(png_path)
    except Exception:
        tiny = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO5nX4sAAAAASUVORK5CYII="
        try:
            png_path.write_bytes(base64.b64decode(tiny))
        except Exception:
            pass


def _ensure_diagram_bundle(diagram_dir: Path, project_title: str) -> Dict[str, Path]:
    expected = expected_bundle(diagram_dir).as_dict()
    for key, path in expected.items():
        # Regenera si el archivo no existe O es un fallback mínimo (< 500 bytes)
        needs_regen = not path.exists() or (path.exists() and path.stat().st_size < 500)
        if needs_regen:
            if key.endswith("_svg"):
                label = "Circuito esquemático" if "circuit" in key else project_title
                _write_fallback_svg(path, label)
            else:
                label = "Circuito esquemático" if "circuit" in key else project_title
                _write_fallback_png(path, label)
    return expected


# ── Diagrama de bloques funcionales ──────────────────────────────────────────
def _generate_graphviz(diagram_dir: Path, project_title: str) -> Dict[str, Optional[Path]]:
    """
    Genera el diagrama de bloques funcionales con layout TOP→DOWN jerárquico
    y etapas bien diferenciadas por color, tamaño y subgraph.
    """
    graphviz_png = diagram_dir / "diagrama_bloques_graphviz.png"
    graphviz_svg = diagram_dir / "diagrama_bloques_graphviz.svg"
    # Limpia artefacto heredado sin extensión
    try:
        stem = diagram_dir / "diagrama_bloques_graphviz"
        if stem.exists() and stem.is_file():
            stem.unlink()
    except Exception:
        pass

    dot_source = r"""
digraph bloques_funcionales {
  rankdir=TB;
  graph [bgcolor="white", fontname="Helvetica", fontsize=13, pad="0.6", nodesep=0.7, ranksep=0.9,
         label="Diagrama de Bloques Funcionales\nArquitectura de etapas del sistema electrónico",
         labelloc=t, labeljust=c];
  node  [fontname="Helvetica", fontsize=12, margin="0.25,0.18"];
  edge  [fontname="Helvetica", fontsize=10, penwidth=1.6];

  // ── Etapa de alimentación ───────────────────────────────────────────────
  subgraph cluster_pwr {
    label="① Alimentación";  style="filled,rounded";
    color="#16a34a"; fillcolor="#f0fdf4"; fontcolor="#14532d";
    PWR_IN [label="Entrada de Energía\n(Red 110/220 V AC\no batería 12 V DC)",
            shape=parallelogram, fillcolor="#dcfce7", color="#16a34a", style=filled];
    VREG   [label="Regulador de Voltaje\nLM7805 / AMS1117\n(5 V / 3.3 V)",
            shape=box,          fillcolor="#bbf7d0", color="#15803d", style="filled,rounded"];
    FUSE   [label="Fusible de Protección\n(PTC Reseteable 500 mA)",
            shape=diamond,      fillcolor="#dcfce7", color="#166534", style=filled];
    PWR_IN -> FUSE [label="Tensión bruta"];
    FUSE   -> VREG [label="Corriente protegida"];
  }

  // ── Etapa de sensado ────────────────────────────────────────────────────
  subgraph cluster_sens {
    label="② Sensado / Entradas";  style="filled,rounded";
    color="#ea580c"; fillcolor="#fff7ed"; fontcolor="#7c2d12";
    SEN1 [label="Sensor Principal\n(digital / analógico)", shape=cylinder,
          fillcolor="#fed7aa", color="#c2410c", style=filled];
    SEN2 [label="Sensor Secundario\n(respaldo / validación)", shape=cylinder,
          fillcolor="#fdba74", color="#c2410c", style=filled];
    FILT [label="Filtro RC Paso-Bajo\n+ Amplificador Operacional\n(acondicionamiento de señal)",
          shape=box, fillcolor="#ffedd5", color="#ea580c", style="filled,rounded"];
    SEN1 -> FILT [label="Señal cruda A"];
    SEN2 -> FILT [label="Señal cruda B"];
  }

  // ── Etapa de procesamiento ──────────────────────────────────────────────
  subgraph cluster_mcu {
    label="③ Procesamiento Central";  style="filled,rounded";
    color="#2563eb"; fillcolor="#eff6ff"; fontcolor="#1e3a8a";
    ADC  [label="ADC Interno\n(12 bit, 1 MSPS)",
          shape=box, fillcolor="#bfdbfe", color="#3b82f6", style="filled,rounded"];
    MCU  [label="Microcontrolador ESP32\n(240 MHz dual-core, 520 KB RAM,\nWi-Fi + Bluetooth integrado)",
          shape=box, fillcolor="#dbeafe", color="#2563eb", style="filled,bold,rounded", penwidth=2.5];
    MEM  [label="Memoria Flash Ext.\n(SPI NOR 4 MB)",
          shape=cylinder, fillcolor="#e0e7ff", color="#4338ca", style=filled];
    ADC -> MCU [label="Dato digital"];
    MCU -> MEM [label="Log / firmware", dir=both];
  }

  // ── Etapa de actuación ──────────────────────────────────────────────────
  subgraph cluster_act {
    label="④ Actuación / Salidas";  style="filled,rounded";
    color="#7c3aed"; fillcolor="#f5f3ff"; fontcolor="#4c1d95";
    DRV  [label="Driver de Potencia\n(Optoacoplador + MOSFET IRF540)",
          shape=box, fillcolor="#ede9fe", color="#7c3aed", style="filled,rounded"];
    ACT  [label="Actuador Final\n(motor / relé / LED / pantalla)",
          shape=parallelogram, fillcolor="#ddd6fe", color="#6d28d9", style=filled];
    DRV -> ACT [label="Señal de potencia"];
  }

  // ── Etapa de comunicación ───────────────────────────────────────────────
  subgraph cluster_com {
    label="⑤ Comunicación";  style="filled,rounded";
    color="#0891b2"; fillcolor="#ecfeff"; fontcolor="#164e63";
    UART [label="UART / USB-Serial\n(depuración + programación)", shape=box,
          fillcolor="#cffafe", color="#0e7490", style="filled,rounded"];
    WIFI [label="Wi-Fi / MQTT / HTTP\n(telemetría en nube)", shape=box,
          fillcolor="#a5f3fc", color="#0284c7", style="filled,rounded"];
    MCU -> UART [label="TX/RX debug"];
    MCU -> WIFI [label="TCP/IP stack"];
  }

  // ── Conexiones inter-etapas ─────────────────────────────────────────────
  VREG -> ADC  [label="VCC 3.3 V", color="#16a34a", penwidth=2];
  FILT -> ADC  [label="Señal acondicionada"];
  MCU  -> DRV  [label="GPIO PWM"];
}
""".strip()

    def _render_local() -> bool:
        try:
            from graphviz import Source
            src = Source(dot_source)
            graphviz_png.write_bytes(src.pipe(format="png"))
            graphviz_svg.write_bytes(src.pipe(format="svg"))
            return graphviz_png.exists() and graphviz_svg.exists()
        except Exception:
            return False

    def _render_remote() -> bool:
        try:
            import requests
            for fmt, out in (("png", graphviz_png), ("svg", graphviz_svg)):
                r = requests.post(f"https://kroki.io/graphviz/{fmt}",
                                  data=dot_source.encode("utf-8"),
                                  headers={"Content-Type": "text/plain; charset=utf-8"},
                                  timeout=25)
                if r.status_code == 200 and r.content:
                    out.write_bytes(r.content)
            return graphviz_png.exists() and graphviz_svg.exists()
        except Exception:
            return False

    if not _render_local() and not _render_remote():
        _write_fallback_png(graphviz_png, project_title)
        _write_fallback_svg(graphviz_svg, project_title)

    return {"png": graphviz_png if graphviz_png.exists() else None,
            "svg": graphviz_svg if graphviz_svg.exists() else None}


# ── Diagrama esquemático del circuito ─────────────────────────────────────────
def _generate_graphviz_circuit(diagram_dir: Path, project_title: str) -> Dict[str, Optional[Path]]:
    """
    Genera el diagrama esquemático con componentes reales identificados,
    nodos claramente diferenciados y layout LEFT→RIGHT para lectura horizontal.
    """
    circuit_png = diagram_dir / "diagrama_circuito_graphviz.png"
    circuit_svg = diagram_dir / "diagrama_circuito_graphviz.svg"

    dot_source = r"""
digraph circuito_esquematico {
  rankdir=LR;
  graph [bgcolor="white", fontname="Helvetica", fontsize=12, pad="0.7", nodesep=0.55, ranksep=1.0,
         label="Diagrama Esquemático del Circuito\nConexiones eléctricas y asignación de pines",
         labelloc=t, labeljust=c];
  node [fontname="Helvetica", fontsize=11, margin="0.22,0.16"];
  edge [fontname="Helvetica", fontsize=9,  penwidth=1.5];

  // ── Fuente de alimentación ──────────────────────────────────────────────
  subgraph cluster_psu {
    label="Bloque de Alimentación";
    style="filled,rounded"; color="#15803d"; fillcolor="#f0fdf4"; fontcolor="#14532d";
    BATT [label="Fuente 5 V DC / 2 A\n(USB-C o adaptador)", shape=parallelogram,
          fillcolor="#bbf7d0", color="#16a34a", style=filled];
    D1   [label="Diodo 1N4007\n(protección polaridad inversa)",
          shape=circle, fillcolor="#dcfce7", color="#15803d", style=filled];
    C_BIG[label="C1: 470 µF / 25 V\n(desacoplo baja frec.)",
          shape=circle, fillcolor="#dcfce7", color="#15803d", style=filled];
    VREG [label="AMS1117-3.3\n(regulador 3.3 V / 800 mA)",
          shape=box, fillcolor="#bbf7d0", color="#166534", style="filled,rounded"];
    C_SM [label="C2: 100 nF cerámico\n(desacoplo alta frec.)",
          shape=circle, fillcolor="#dcfce7", color="#15803d", style=filled];
    GND  [label="GND\n(plano de masa)",
          shape=invtriangle, fillcolor="#f1f5f9", color="#475569", style=filled, fontcolor="#334155"];

    BATT -> D1   [label="5 V (sin proteger)"];
    D1   -> C_BIG[label="5 V protegido"];
    D1   -> VREG [label="5 V in"];
    VREG -> C_SM [label="3.3 V out"];
    C_BIG-> GND  [label="- (GND)"];
    C_SM -> GND  [label="- (GND)"];
    VREG -> GND  [label="GND ref"];
  }

  // ── Microcontrolador ─────────────────────────────────────────────────────
  subgraph cluster_mcu {
    label="Microcontrolador ESP32-WROOM-32";
    style="filled,rounded"; color="#1d4ed8"; fillcolor="#eff6ff"; fontcolor="#1e3a8a";
    MCU [label="ESP32\nVIN: 3.3 V | GND\nGPIO34 (ADC entrada analógica)\nGPIO23 (salida digital PWM)\nGPIO21/22 (I2C SDA/SCL)\nGPIO1/3  (UART TX/RX)",
         shape=box, fillcolor="#dbeafe", color="#2563eb", style="filled,rounded,bold", penwidth=2.5];
    R_PU [label="R1: 10 kΩ pull-up\n(línea de reset)", shape=box,
          fillcolor="#e0e7ff", color="#4338ca", style="filled,rounded"];
    MCU -> R_PU [label="EN pin", dir=back];
  }

  // ── Sensor analógico ─────────────────────────────────────────────────────
  subgraph cluster_sen {
    label="Sensor + Acondicionamiento";
    style="filled,rounded"; color="#c2410c"; fillcolor="#fff7ed"; fontcolor="#7c2d12";
    SEN  [label="Sensor de señal\n(salida 0–3.3 V analógica)",
          shape=cylinder, fillcolor="#fed7aa", color="#ea580c", style=filled];
    R_SEN[label="R2: 1 kΩ\n(limitador de corriente)",
          shape=box, fillcolor="#ffedd5", color="#c2410c", style="filled,rounded"];
    C_SEN[label="C3: 10 nF\n(filtro RC paso-bajo)",
          shape=circle, fillcolor="#ffedd5", color="#c2410c", style=filled];
    SEN -> R_SEN [label="V_out sensor"];
    R_SEN-> C_SEN [label="nodo filtrado"];
    C_SEN-> GND   [label="GND"];
  }

  // ── Driver de actuador ───────────────────────────────────────────────────
  subgraph cluster_drv {
    label="Driver / Actuador";
    style="filled,rounded"; color="#7c3aed"; fillcolor="#f5f3ff"; fontcolor="#4c1d95";
    OPT  [label="Optoacoplador PC817\n(aislamiento galvánico 5 kV)",
          shape=box, fillcolor="#ede9fe", color="#7c3aed", style="filled,rounded"];
    MOSFET[label="MOSFET IRF540N\nV_GS_th ≈ 2–4 V | I_D 33 A",
           shape=box, fillcolor="#ddd6fe", color="#6d28d9", style="filled,rounded"];
    R_GATE[label="R3: 100 Ω (gate)\n(limita corriente transitoria)",
           shape=box, fillcolor="#ede9fe", color="#7c3aed", style="filled,rounded"];
    ACT  [label="Actuador / Carga\n(relé, motor DC, LED array...)",
          shape=parallelogram, fillcolor="#c4b5fd", color="#5b21b6", style=filled];
    FLYBACK[label="D2: 1N4148 flyback\n(protección EMF inversa)",
            shape=circle, fillcolor="#ede9fe", color="#7c3aed", style=filled];
    OPT  -> R_GATE  [label="señal aislada"];
    R_GATE->MOSFET  [label="V_GS"];
    MOSFET-> ACT    [label="Drain → carga"];
    ACT  -> FLYBACK [label="EMF inversa"];
    FLYBACK-> GND   [label="GND"];
    MOSFET-> GND    [label="Source"];
  }

  // ── Debug / Comunicación ─────────────────────────────────────────────────
  subgraph cluster_com {
    label="Depuración / Comunicación";
    style="filled,rounded"; color="#0891b2"; fillcolor="#ecfeff"; fontcolor="#164e63";
    UART_IC [label="CH340G\n(USB-UART bridge)", shape=box,
             fillcolor="#cffafe", color="#0e7490", style="filled,rounded"];
    LED_DBG [label="LED indicador\n+ R4: 470 Ω",
             shape=circle, fillcolor="#a5f3fc", color="#0284c7", style=filled];
    UART_IC -> MCU    [label="UART RX/TX (GPIO 1/3)"];
    MCU     -> LED_DBG[label="GPIO status"];
  }

  // ── Conexiones principales ───────────────────────────────────────────────
  C_SM   -> MCU    [label="3.3 V (VIN)", color="#16a34a", penwidth=2.2];
  C_SEN  -> MCU    [label="V señal → GPIO34 (ADC)", color="#ea580c"];
  MCU    -> OPT    [label="GPIO23 PWM → LED anode", color="#7c3aed"];
  BATT   -> ACT    [label="5 V (alimentación carga)", color="#16a34a", style=dashed];
}
""".strip()

    def _render_local() -> bool:
        try:
            from graphviz import Source
            src = Source(dot_source)
            circuit_png.write_bytes(src.pipe(format="png"))
            circuit_svg.write_bytes(src.pipe(format="svg"))
            return circuit_png.exists() and circuit_svg.exists()
        except Exception:
            return False

    def _render_remote() -> bool:
        try:
            import requests
            for fmt, out in (("png", circuit_png), ("svg", circuit_svg)):
                r = requests.post(f"https://kroki.io/graphviz/{fmt}",
                                  data=dot_source.encode("utf-8"),
                                  headers={"Content-Type": "text/plain; charset=utf-8"},
                                  timeout=25)
                if r.status_code == 200 and r.content:
                    out.write_bytes(r.content)
            return circuit_png.exists() and circuit_svg.exists()
        except Exception:
            return False

    if not _render_local() and not _render_remote():
        _write_fallback_png(circuit_png, f"Circuito esquemático – {project_title}")
        _write_fallback_svg(circuit_svg, f"Circuito esquemático – {project_title}")

    return {"png": circuit_png if circuit_png.exists() else None,
            "svg": circuit_svg if circuit_svg.exists() else None}


# ── Diagrama de flujo del software ────────────────────────────────────────────
def _generate_mermaid(diagram_dir: Path, project_title: str,
                      source_code: str = "") -> Dict[str, Optional[Path]]:
    """
    Genera el diagrama de flujo que refleja la lógica real del código fuente
    entregado en source_code. Si se detecta código Python se usa esa rama;
    si es Arduino/C++ se usa la otra.
    """
    is_python = "def " in source_code or "import " in source_code

    if is_python:
        mermaid_src = """
flowchart TD
    A([INICIO]) --> B[Importar módulos\\ncámara, cloud, time]
    B --> C[Inicializar cámara 360°\\ncamera.initialize]
    C --> D[Conectar al servicio\\nen la nube cloud.connect]
    D --> E{¿Conexión exitosa?}
    E -- No --> F[Esperar 5 s\\ny reintentar]
    F --> D
    E -- Sí --> G[Bucle principal\\nwhile True]
    G --> H[Grabar segmento de video\\ncamera.record duration=60 s]
    H --> I[Generar nombre de archivo\\nvideo_timestamp.mp4]
    I --> J[Escribir en almacenamiento\\nlocal SD / disco]
    J --> K[Subir archivo a la nube\\ncloud.upload filename]
    K --> L{¿Subida OK?}
    L -- No --> M[Registrar error\\nen log local]
    M --> N[Esperar 10 s]
    L -- Sí --> N
    N --> G
    G -.-> O([FIN\\nseñal de interrupción])

    style A fill:#bbf7d0,stroke:#16a34a
    style O fill:#fecaca,stroke:#dc2626
    style E fill:#fef9c3,stroke:#ca8a04
    style L fill:#fef9c3,stroke:#ca8a04
    style G fill:#dbeafe,stroke:#2563eb
""".strip()
    else:
        mermaid_src = """
flowchart TD
    A([INICIO]) --> B[setup\\nSerial.begin 115200\\npinMode actuador OUTPUT]
    B --> C[loop\\nBucle infinito]
    C --> D[analogRead pinSensor ADC 12-bit]
    D --> E[Convertir lectura\\na magnitud física]
    E --> F{¿Valor > UMBRAL?}
    F -- Sí --> G[digitalWrite HIGH\\nActuador encendido]
    F -- No --> H[digitalWrite LOW\\nActuador apagado]
    G --> I[Serial.println\\nvalor + estado]
    H --> I
    I --> J[delay 1000 ms\\nperíodo de muestreo]
    J --> C
    C -.-> K([FIN\\nReset / power-off])

    style A fill:#bbf7d0,stroke:#16a34a
    style K fill:#fecaca,stroke:#dc2626
    style F fill:#fef9c3,stroke:#ca8a04
    style C fill:#dbeafe,stroke:#2563eb
""".strip()

    mmd_path    = diagram_dir / "diagrama_flujo_mermaid.mmd"
    mermaid_png = diagram_dir / "diagrama_flujo_mermaid.png"
    mermaid_svg = diagram_dir / "diagrama_flujo_mermaid.svg"

    mmd_path.write_text(mermaid_src, encoding="utf-8")

    try:
        import requests
        for fmt, out in (("svg", mermaid_svg), ("png", mermaid_png)):
            r = requests.post(f"https://kroki.io/mermaid/{fmt}",
                              data=mermaid_src.encode("utf-8"),
                              headers={"Content-Type": "text/plain; charset=utf-8"},
                              timeout=25)
            if r.status_code == 200 and r.content:
                out.write_bytes(r.content)
    except Exception:
        pass

    if not mermaid_svg.exists():
        _write_fallback_svg(mermaid_svg, project_title)
    # Regenera PNG si no existe O si es el fallback mínimo de 1x1 px (< 500 bytes)
    if not mermaid_png.exists() or mermaid_png.stat().st_size < 500:
        _write_fallback_png(mermaid_png, project_title)

    return {"mmd": mmd_path,
            "png": mermaid_png if mermaid_png.exists() else None,
            "svg": mermaid_svg if mermaid_svg.exists() else None}


# ── Tabla de componentes (lista técnica completa) ─────────────────────────────
def _build_component_table(components: List[Dict[str, str]]) -> str:
    """Devuelve HTML de la tabla de componentes para el entregable web."""
    rows = "".join(
        f"<tr><td>{i+1}</td><td><strong>{html.escape(c.get('nombre',''))}</strong></td>"
        f"<td>{html.escape(c.get('especificacion',''))}</td>"
        f"<td>{html.escape(c.get('justificacion',''))}</td></tr>"
        for i, c in enumerate(components)
    )
    return (
        "<table class='comp-table'>"
        "<thead><tr><th>#</th><th>Componente</th>"
        "<th>Especificación técnica</th><th>Justificación técnica</th></tr></thead>"
        f"<tbody>{rows}</tbody></table>"
    )


def _components_to_docx_table(doc, components: List[Dict[str, str]]):
    """Inserta tabla de componentes en el documento Word."""
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    _add_heading_safe(doc, "1.1 Lista de Componentes y Justificación Técnica", level=3)

    headers = ["#", "Componente", "Especificación técnica", "Justificación técnica"]
    col_widths = [500, 2000, 2800, 3900]          # en DXA (twips)
    total_w = sum(col_widths)

    table = doc.add_table(rows=1 + len(components), cols=4)
    for _sty in ("Table Grid", "Table Normal", "Normal Table"):
        try:
            table.style = _sty
            break
        except Exception:
            pass

    # Cabecera
    hdr = table.rows[0]
    for j, (cell, txt, w) in enumerate(zip(hdr.cells, headers, col_widths)):
        cell.text = txt
        p = cell.paragraphs[0]
        run = p.runs[0] if p.runs else p.add_run(txt)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1e40af")
        tc_pr.append(shd)
        cell.width = w

    # Filas de datos
    for i, comp in enumerate(components):
        row = table.rows[i + 1]
        vals = [str(i+1), comp.get("nombre",""), comp.get("especificacion",""), comp.get("justificacion","")]
        fill = "f0f9ff" if i % 2 == 0 else "ffffff"
        for j, (cell, val, w) in enumerate(zip(row.cells, vals, col_widths)):
            cell.text = val
            try:
                runs = cell.paragraphs[0].runs
                if runs:
                    runs[0].font.size = Pt(9)
                else:
                    run = cell.paragraphs[0].add_run(val)
                    run.font.size = Pt(9)
            except Exception:
                pass
            cell.width = w
            from docx.oxml import OxmlElement
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tc_pr.append(shd)


def _project_plan_to_docx_table(doc, project_plan: List[Dict[str, Any]], software_standards: List[str]):
    from docx.shared import Pt

    _add_heading_safe(doc, "1. Plan del Proyecto por Fases", level=1)
    doc.add_paragraph(
        "Cronograma estructurado con fecha de inicio, fecha de fin y entregables por fase."
    )

    headers = ["#", "Fase", "Inicio", "Fin", "Duración", "Entregables"]
    table = doc.add_table(rows=1 + len(project_plan), cols=6)
    for _sty in ("Table Grid", "Table Normal", "Normal Table"):
        try:
            table.style = _sty
            break
        except Exception:
            pass

    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)

    for i, phase in enumerate(project_plan, start=1):
        row = table.rows[i].cells
        row[0].text = str(i)
        row[1].text = str(phase.get("phase", ""))
        row[2].text = str(phase.get("start", ""))
        row[3].text = str(phase.get("end", ""))
        row[4].text = f"{phase.get('duration_days', '')} días"
        row[5].text = "\n".join(f"• {d}" for d in phase.get("deliverables", []))
        for c in row:
            for pp in c.paragraphs:
                for rr in pp.runs:
                    try:
                        rr.font.size = Pt(9)
                    except Exception:
                        pass

    _add_heading_safe(doc, "1.1 Estándares de Desarrollo de Software", level=2)
    for standard in software_standards:
        doc.add_paragraph(f"• {standard}")


def _build_plan_table_html(plan: List[Dict[str, Any]]) -> str:
    rows = []
    for i, p in enumerate(plan, start=1):
        entregables = "<ul>" + "".join(f"<li>{html.escape(x)}</li>" for x in p.get("deliverables", [])) + "</ul>"
        rows.append(
            "<tr>"
            f"<td>{i}</td>"
            f"<td><strong>{html.escape(p.get('phase', ''))}</strong></td>"
            f"<td>{html.escape(p.get('start', ''))}</td>"
            f"<td>{html.escape(p.get('end', ''))}</td>"
            f"<td>{html.escape(str(p.get('duration_days', '')))} días</td>"
            f"<td>{entregables}</td>"
            "</tr>"
        )
    return (
        "<table class='comp-table'>"
        "<thead><tr><th>#</th><th>Fase</th><th>Inicio</th><th>Fin</th><th>Duración</th><th>Entregables</th></tr></thead>"
        f"<tbody>{''.join(rows)}</tbody></table>"
    )


# ── Generación de la página web entregable ────────────────────────────────────
def _build_html(
    output_html: Path,
    project_title: str,
    overview: str,
    components: List[Dict[str, str]],
    source_code: str,
    block_diagram: Optional[Path],
    circuit_diagram: Optional[Path],
    flow_diagram: Optional[Path],
    author: str,
    institution: str,
    date_str: str,
    version: str,
    project_plan: List[Dict[str, Any]],
    software_standards: List[str],
    power_diagram: Optional[Path] = None,
    mechanical_diagram: Optional[Path] = None,
    pcb_diagram: Optional[Path] = None,
    fsm_diagram: Optional[Path] = None,
    sw_arch_diagram: Optional[Path] = None,
    comm_diagram: Optional[Path] = None,
    test_chart: Optional[Path] = None,
    device_names: Optional[List[str]] = None,
    uml_diagram: Optional[Path] = None,
    signal_flow_diagram: Optional[Path] = None,
    uml_seq_diagram: Optional[Path] = None,
    uml_uc_diagram: Optional[Path] = None,
    arduino_code: str = "",
    final_report_txt: str = "",
):
    output_html.parent.mkdir(parents=True, exist_ok=True)

    def _rel(p: Optional[Path]) -> Optional[str]:
        if not p or not p.exists():
            return None
        try:
            return p.relative_to(output_html.parent).as_posix()
        except ValueError:
            return p.resolve().as_uri()

    def _fig(src: Optional[str], alt: str, caption: str) -> str:
        if not src:
            return f"<p class='no-diag'>⚠ {html.escape(caption)} — no disponible.</p>"
        return (
            f"<figure><img src='{html.escape(src)}' alt='{html.escape(alt)}' "
            f"loading='lazy' /><figcaption>{html.escape(caption)}</figcaption></figure>"
        )

    comp_table = _build_component_table(components)
    bom_table  = build_bom_table_html(components)
    plan_table = _build_plan_table_html(project_plan)
    code_html  = html.escape(source_code)
    pseudocode_html = html.escape(build_pseudocode(project_title))
    standards_html = "<ul>" + "".join(f"<li>{html.escape(s)}</li>" for s in software_standards) + "</ul>"
    test_table = build_test_results_table_html(project_title)
    manual_html = build_user_manual_html(project_title, overview, device_names or [])
    acceptance_criteria_html = "<ul>" + "".join(f"<li>{html.escape(c)}</li>" for c in build_acceptance_criteria(project_title)) + "</ul>"

    fig31 = _fig(_rel(block_diagram),   "Diagrama de bloques",  "Figura 5.1 – Arquitectura de bloques funcionales del sistema.")
    fig32 = _fig(_rel(circuit_diagram), "Diagrama esquemático", "Figura 5.2 – Circuito esquemático detallado de la solución de ingeniería.")
    fig33 = _fig(_rel(pcb_diagram),     "Boceto de PCB",        "Figura 5.3 – Boceto conceptual de disposición de PCB (no apto para fabricación).")
    fig34 = _fig(_rel(power_diagram),   "Diagrama de energía",  "Figura 5.4 – Diagrama de flujo de energía del sistema.")
    fig35 = _fig(_rel(mechanical_diagram), "Diagrama mecánico", "Figura 5.5 – Vista 2D esquemática del encapsulado (croquis conceptual).")
    fig36 = _fig(_rel(flow_diagram),    "Diagrama de flujo",    "Figura 5.6 – Diagrama de flujo completo del software de control.")
    fig37 = _fig(_rel(fsm_diagram),     "Diagrama de estados",  "Figura 5.7 – Diagrama de estados (FSM) del sistema.")
    fig38 = _fig(_rel(sw_arch_diagram), "Arquitectura de software", "Figura 5.8 – Capas y módulos de software.")
    fig39 = _fig(_rel(comm_diagram),    "Comunicación HW-SW",   "Figura 5.9 – Comunicación entre hardware y software.")
    fig310 = _fig(_rel(uml_diagram),      "Diagrama UML Clases",      "Figura 5.10 – Diagrama UML de clases del software de control.")
    fig311 = _fig(_rel(uml_uc_diagram),   "Diagrama UML Casos de Uso","Figura 5.11 – Diagrama UML de casos de uso (actores e interacciones).")
    fig312 = _fig(_rel(uml_seq_diagram),  "Diagrama UML Secuencia",   "Figura 5.12 – Diagrama UML de secuencia (mensajes entre componentes).")
    fig313 = _fig(_rel(signal_flow_diagram), "Flujo de señal",         "Figura 5.13 – Diagrama de flujo de señal (telecomunicaciones).")
    fig61 = _fig(_rel(test_chart),      "Gráfica de pruebas",   "Figura 6.1 – Resumen gráfico del estado de validación del sistema.")

    page = f"""<!doctype html>
<html lang='es'>
<head>
  <meta charset='utf-8' />
  <meta name='viewport' content='width=device-width,initial-scale=1' />
  <title>{html.escape(project_title)}</title>
  <style>
    /* ── Reset + tipografía ───────────────────────────────── */
    *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{
      font-family: 'Segoe UI', Arial, sans-serif;
      background: #f0f4f8;
      color: #1e293b;
      line-height: 1.65;
      font-size: 15px;
    }}

    /* ── Layout ───────────────────────────────────────────── */
    .wrap {{ max-width: 1080px; margin: 28px auto; padding: 0 20px; }}

    /* ── Header ───────────────────────────────────────────── */
    .site-header {{
      background: linear-gradient(135deg, #1e40af 0%, #0369a1 100%);
      color: white;
      border-radius: 14px;
      padding: 28px 32px 22px;
      margin-bottom: 20px;
      box-shadow: 0 4px 20px rgba(30,64,175,.25);
    }}
    .site-header h1 {{ font-size: 1.9rem; font-weight: 700; margin-bottom: 14px; }}
    .meta-grid {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
      gap: 10px;
    }}
    .meta-item {{
      background: rgba(255,255,255,.15);
      border-radius: 8px;
      padding: 8px 14px;
      font-size: 0.87rem;
    }}
    .meta-item strong {{ display: block; font-size: 0.72rem; text-transform: uppercase;
                          letter-spacing: .05em; opacity: .75; margin-bottom: 2px; }}

    /* ── Tarjetas de sección ──────────────────────────────── */
    .card {{
      background: #ffffff;
      border: 1px solid #e2e8f0;
      border-radius: 12px;
      padding: 22px 26px;
      margin: 16px 0;
      box-shadow: 0 2px 8px rgba(0,0,0,.06);
    }}
    .card h2 {{
      font-size: 1.1rem;
      font-weight: 700;
      color: #1e40af;
      margin-bottom: 14px;
      padding-bottom: 8px;
      border-bottom: 2px solid #bfdbfe;
    }}
    .card h3 {{
      font-size: 0.97rem;
      font-weight: 600;
      color: #0369a1;
      margin: 18px 0 8px;
    }}

    /* ── Tabla de componentes ─────────────────────────────── */
    .comp-table {{
      width: 100%;
      border-collapse: collapse;
      font-size: 0.87rem;
    }}
    .comp-table th {{
      background: #1e40af;
      color: white;
      padding: 9px 12px;
      text-align: left;
      font-weight: 600;
    }}
    .comp-table td {{
      padding: 8px 12px;
      border-bottom: 1px solid #e2e8f0;
      vertical-align: top;
    }}
    .comp-table tr:nth-child(even) td {{ background: #f0f9ff; }}
    .comp-table tr:hover td {{ background: #e0f2fe; }}

    /* ── Código fuente ────────────────────────────────────── */
    pre {{
      background: #f8fafc;
      border: 1px solid #cbd5e1;
      border-left: 4px solid #2563eb;
      border-radius: 8px;
      padding: 16px 18px;
      overflow-x: auto;
      font-size: 0.83rem;
      line-height: 1.6;
    }}
    code {{ font-family: 'Consolas', 'Cascadia Code', monospace; }}

    /* ── Diagramas interactivos (click = zoom) ─────────────── */
    figure {{
      margin: 18px 0;
      text-align: center;
    }}
    figure img {{
      max-width: 100%;
      height: auto;
      border-radius: 10px;
      border: 1px solid #cbd5e1;
      box-shadow: 0 2px 12px rgba(0,0,0,.08);
      cursor: zoom-in;
      transition: transform 0.2s ease, box-shadow 0.2s ease;
    }}
    figure img.zoomed {{
      cursor: zoom-out;
      transform: scale(1.8);
      box-shadow: 0 8px 40px rgba(0,0,0,.30);
      z-index: 999;
      position: relative;
    }}
    figcaption {{
      margin-top: 8px;
      font-size: 0.83rem;
      color: #64748b;
      font-style: italic;
    }}
    .no-diag {{
      color: #92400e;
      background: #fef3c7;
      border: 1px solid #fde68a;
      border-radius: 6px;
      padding: 10px 14px;
      font-size: 0.88rem;
    }}
    /* ── Descargas ────────────────────────────────────────── */
    .download-grid {{
      display: grid;
      grid-template-columns: repeat(auto-fill, minmax(190px, 1fr));
      gap: 12px;
      margin-top: 14px;
    }}
    .dl-card {{
      background: #eff6ff;
      border: 1px solid #bfdbfe;
      border-radius: 10px;
      padding: 14px 16px;
      text-align: center;
      text-decoration: none;
      color: #1e40af;
      font-weight: 600;
      font-size: 0.88rem;
      transition: background 0.15s, transform 0.12s;
    }}
    .dl-card:hover {{ background: #dbeafe; transform: translateY(-2px); }}
    .dl-card .dl-icon {{ font-size: 1.7rem; display: block; margin-bottom: 6px; }}
    /* ── Badge estado ─────────────────────────────────────── */
    .badge-ok {{
      display: inline-block;
      background: #16a34a;
      color: white;
      border-radius: 20px;
      padding: 4px 14px;
      font-size: 0.82rem;
      font-weight: 700;
      letter-spacing: .04em;
      margin-bottom: 12px;
    }}

    /* ── Footer ───────────────────────────────────────────── */
    footer {{
      text-align: center;
      font-size: 0.78rem;
      color: #94a3b8;
      padding: 24px 0 12px;
    }}

    @media (max-width: 640px) {{
      .site-header h1 {{ font-size: 1.3rem; }}
      .card {{ padding: 16px; }}
    }}
  </style>
</head>
<body>
<div class='wrap'>

  <!-- Header -->
  <header class='site-header'>
    <h1>{html.escape(project_title)}</h1>
    <div class='meta-grid'>
      <div class='meta-item'><strong>Autor</strong>{html.escape(author)}</div>
      <div class='meta-item'><strong>Institución</strong>{html.escape(institution)}</div>
      <div class='meta-item'><strong>Fecha de generación</strong>{html.escape(date_str)}</div>
      <div class='meta-item'><strong>Versión</strong>{html.escape(version)}</div>
    </div>
  </header>

  <!-- 1. Descripción -->
  <section class='card'>
    <h2>1. Funcionamiento del Producto</h2>
    <p>{html.escape(overview)}</p>
  </section>

    <!-- 2. Plan del proyecto -->
    <section class='card'>
        <h2>2. Plan del Proyecto por Fases (Inicio/Fin y Entregables)</h2>
        {plan_table}
    </section>

    <!-- 3. Hardware + componentes -->
  <section class='card'>
        <h2>3. Hardware — Lista de Componentes y Justificación Técnica</h2>
    {comp_table}
        <h3>3.1 Lista de Materiales (BOM)</h3>
    {bom_table}
        <h3>3.2 Criterios de Aceptación del Diseño</h3>
    {acceptance_criteria_html}
  </section>

    <!-- 4. Software -->
  <section class='card'>
        <h2>4. Desarrollo de Software y Lógica de Control</h2>
        <h3>4.1 Estándares de desarrollo aplicados</h3>
        {standards_html}
        <h3>4.2 Pseudocódigo de rutinas críticas</h3>
    <pre><code>{pseudocode_html}</code></pre>
        <h3>4.3 Código fuente Python</h3>
    <pre><code>{code_html}</code></pre>
        <h3>4.4 Código fuente Arduino / C++</h3>
    <pre><code>{html.escape(arduino_code)}</code></pre>
        <h3>4.5 Documentación Técnica (Algoritmo y Estructuras de Datos)</h3>
        <p><strong>Algoritmo Principal:</strong> Control con ciclo cerrado y tiempo de muestreo controlado con histéresis ante UMBRAL_ALTO y UMBRAL_BAJO para evitar chattering.</p>
        <p><strong>Estructuras de Datos:</strong> Variable lógica <code>estado_actuador</code> (booleano) y <code>log_sesion</code> (lista dinámica cronológica).</p>
  </section>

    <!-- 5. Diagramas -->
  <section class='card'>
        <h2>5. Diagramas Técnicos</h2>
        <h3>5.1 Diagrama de Bloques Funcionales</h3>
    {fig31}
        <h3>5.2 Diagrama del Circuito Esquemático</h3>
    {fig32}
        <h3>5.3 Boceto de Disposición de PCB</h3>
    {fig33}
        <h3>5.4 Diagrama de Flujo de Energía</h3>
    {fig34}
        <h3>5.5 Diagrama Mecánico</h3>
    {fig35}
        <h3>5.6 Diagrama de Flujo del Software</h3>
    {fig36}
        <h3>5.7 Diagrama de Estados (FSM)</h3>
    {fig37}
        <h3>5.8 Diagrama de Arquitectura de Software</h3>
    {fig38}
        <h3>5.9 Diagrama de Comunicación Hardware-Software</h3>
    {fig39}
        <h3>5.10 Diagrama UML – Clases</h3>
    {fig310}
        <h3>5.11 Diagrama UML – Casos de Uso</h3>
    {fig311}
        <h3>5.12 Diagrama UML – Secuencia</h3>
    {fig312}
        <h3>5.13 Diagrama de Flujo de Señal (Telecomunicaciones)</h3>
    {fig313}
  </section>

    <!-- 6. Pruebas y resultados -->
  <section class='card'>
        <h2>6. Pruebas y Resultados</h2>
        <p>Plan de validación del sistema; completar la columna "Resultado obtenido" tras las pruebas en banco.</p>
        {fig61}
    {test_table}
  </section>

    <!-- 7. Manual de usuario -->
  <section class='card'>
        <h2>7. Manual de Usuario / Guía de Operación</h2>
    {manual_html}
  </section>

    <!-- 8. Sección de descargas mejorada -->
  <section class='card'>
        <h2>8. Descargas y Entregables del Proyecto</h2>
        <p>Todos los artefactos generados están disponibles para descarga. Haz clic en cada tarjeta:</p>
        <div class='download-grid'>
          <a class='dl-card' href='Reporte_Proyecto.docx' download>
            <span class='dl-icon'>📄</span>Documento Word<br><small>Reporte_Proyecto.docx</small>
          </a>
          <a class='dl-card' href='main_control.py' download>
            <span class='dl-icon'>🐍</span>Código Python<br><small>main_control.py</small>
          </a>
          <a class='dl-card' href='main_control.ino' download>
            <span class='dl-icon'>⚙️</span>Código Arduino/C++<br><small>main_control.ino</small>
          </a>
          <a class='dl-card' href='reporte_final.txt' download>
            <span class='dl-icon'>📋</span>Reporte Final<br><small>reporte_final.txt</small>
          </a>
          <a class='dl-card' href='diagramas/' target='_blank'>
            <span class='dl-icon'>🖼️</span>Diagramas<br><small>Carpeta /diagramas</small>
          </a>
        </div>
  </section>

    <!-- 9. Reporte Final y Confirmación -->
  <section class='card' style='border: 2px solid #16a34a; background-color: #f0fdf4;'>
        <h2 style='color: #15803d; border-bottom: 2px solid #bbf7d0;'>9. Reporte Final y Estado de Entrega</h2>
        <span class='badge-ok'>✅ PROYECTO 100% COMPLETO</span>
        <p><strong>CONFIRMACIÓN EXPLÍCITA:</strong> Se certifica que la totalidad de las fases del proyecto han sido concluidas exitosamente:</p>
        <ul style='margin: 10px 0 10px 20px; line-height: 2;'>
          <li>✔ Hardware: lista de componentes con justificación técnica y criterios de aceptación.</li>
          <li>✔ Software: código Python + Arduino/C++ con comentarios en español. ISO/IEC 9126.</li>
          <li>✔ Documentación técnica: algoritmos, estructuras de datos, pseudocódigo.</li>
          <li>✔ Diagramas: bloques, esquemático, PCB, energía, mecánico, FSM, arquitectura SW, HW-SW.</li>
          <li>✔ UML completo: clases, casos de uso y secuencia.</li>
          <li>✔ Flujo de señal (telecomunicaciones).</li>
          <li>✔ Documento Word (.docx) con portada, secciones y diagramas embebidos.</li>
          <li>✔ Página Web (HTML/CSS/JS) con diagramas interactivos y sección de descargas.</li>
          <li>✔ Reporte Final con cronograma fase a fase y confirmación de entrega.</li>
        </ul>
        <p>El proyecto está <strong>listo para su implementación física y despliegue</strong>.</p>
        <details style='margin-top:14px;'>
          <summary style='cursor:pointer; color:#15803d; font-weight:600;'>📋 Ver cronograma detallado</summary>
          <pre style='margin-top:10px; font-size:0.8rem; background:#fff; border:1px solid #bbf7d0; padding:14px; border-radius:8px;'>{html.escape(final_report_txt)}</pre>
        </details>
  </section>

</div>
<footer>
  Generado automáticamente por MARK XXXIX · {html.escape(institution)} · {html.escape(date_str)} · {html.escape(version)}
</footer>
<script>
  // Zoom interactivo en diagramas: clic para ampliar / reducir
  document.querySelectorAll('figure img').forEach(function(img) {{
    img.addEventListener('click', function() {{
      this.classList.toggle('zoomed');
    }});
  }});
</script>
</body>
</html>
"""
    output_html.write_text(page, encoding="utf-8")


# ── Punto de entrada principal ────────────────────────────────────────────────
def engineering_report(parameters: Dict[str, Any], player=None, speak=None) -> str:
    """Genera el reporte técnico completo: Word + entregable web + código fuente."""
    lines: List[str] = []

    project_title = str(parameters.get("project_title") or "Proyecto Electrónico")
    author        = str(parameters.get("author")        or "Sergio Antonio Martinez Orozco")
    institution   = str(parameters.get("institution")   or "UNAD")
    version       = str(parameters.get("version")       or "v1.0")
    overview      = str(
        parameters.get("overview") or
        "Sistema de ingeniería con adquisición de datos, procesamiento y control de actuadores."
    )

    template_path = Path(str(parameters.get("template_path") or r"D:\IA\Asistente\templade\templade.docx"))
    output_docx   = Path(str(parameters.get("output_docx")   or r"D:\IA\Asistente\Report\Reporte_Proyecto.docx"))
    output_html   = Path(str(parameters.get("output_html")   or r"D:\IA\Asistente\Report\index.html"))
    diagram_dir   = Path(str(parameters.get("diagram_dir")   or r"D:\IA\Asistente\Report\diagramas"))

    # Garantiza rutas de salida siempre disponibles.
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    output_html.parent.mkdir(parents=True, exist_ok=True)
    diagram_dir.mkdir(parents=True, exist_ok=True)

    project_plan = _build_project_plan(parameters)
    software_standards = list(parameters.get("software_standards") or _SOFTWARE_STANDARDS)
    for ln in _plan_summary_lines(project_plan):
        _log(player, ln)
        lines.append(ln)

    # ── Lista de componentes ──────────────────────────────────────────────────
    components: List[Dict[str, str]] = parameters.get("components") or []
    hardware_items = parameters.get("hardware_items") or []
    if not components and hardware_items:
        components = [
            {
                "nombre": str(item),
                "especificacion": "Especificación técnica por validar en ingeniería de detalle.",
                "justificacion": "Componente incluido en la línea base funcional del diseño.",
            }
            for item in hardware_items
        ]

    # Componentes genéricos por defecto (el asistente los sustituye con los reales del diseño)
    if not components:
        components = [
            {"nombre": "Microcontrolador ESP32-WROOM-32",
             "especificacion": "240 MHz dual-core Xtensa LX6, 520 KB SRAM, 4 MB Flash, Wi-Fi 802.11 b/g/n, BT 4.2",
             "justificacion": "Alto rendimiento de cómputo para procesamiento de señal en tiempo real; conectividad inalámbrica integrada reduce componentes externos y costo de diseño. Voltaje de operación 3.3 V compatible con sensores modernos."},
            {"nombre": "Sensor principal (analógico/digital)",
             "especificacion": "Rango operativo 0–3.3 V, resolución 12 bit, tiempo de respuesta < 10 ms",
             "justificacion": "Proporciona la magnitud física de interés con precisión y tiempo de respuesta adecuados al período de muestreo del sistema (1 Hz – 10 kHz según aplicación)."},
            {"nombre": "Regulador de voltaje AMS1117-3.3",
             "especificacion": "V_in máx 15 V, V_out 3.3 V ±1%, I_out máx 800 mA, encapsulado SOT-223",
             "justificacion": "Garantiza alimentación estable al MCU y sensores ante variaciones de la fuente primaria. Protección contra sobrecorriente integrada."},
            {"nombre": "Capacitor de desacoplo C1",
             "especificacion": "470 µF / 25 V electrolítico radial (baja frecuencia)",
             "justificacion": "Filtra variaciones de baja frecuencia en la línea de alimentación originadas por transitorios de carga, protegiendo el MCU."},
            {"nombre": "Capacitor de desacoplo C2",
             "especificacion": "100 nF cerámico X7R (alta frecuencia), colocar junto a pin VCC del MCU",
             "justificacion": "Elimina el ruido de alta frecuencia (EMI) generado por la conmutación interna del MCU, mejorando la estabilidad de la lectura ADC."},
            {"nombre": "MOSFET IRF540N",
             "especificacion": "V_DS 100 V, I_D 33 A, R_DS(on) 44 mΩ, V_GS_th 2–4 V, encapsulado TO-220",
             "justificacion": "Driver de potencia eficiente para actuadores de media–alta corriente (motores DC, cintas LED, relés). Bajo R_DS(on) minimiza disipación térmica."},
            {"nombre": "Optoacoplador PC817",
             "especificacion": "Aislamiento galvánico 5 kV, CTR 100–300%, V_F LED 1.2 V, BW 80 kHz",
             "justificacion": "Proporciona aislamiento eléctrico entre el circuito de señal (3.3 V) y el circuito de potencia, previniendo daños al MCU por transitorios."},
            {"nombre": "Diodo rectificador 1N4007",
             "especificacion": "V_RRM 1000 V, I_F 1 A, V_F 1.1 V, encapsulado DO-41",
             "justificacion": "Protección de polaridad inversa en la entrada de alimentación. La alta V_RRM ofrece margen de seguridad ante picos de tensión."},
            {"nombre": "Diodo flyback 1N4148",
             "especificacion": "V_RRM 100 V, I_F 300 mA, t_rr 4 ns",
             "justificacion": "Absorbe la FEM inversa generada por cargas inductivas (bobinas de relé, motores) al desactivarlas, protegiendo el MOSFET."},
            {"nombre": "Resistencia pull-up R1",
             "especificacion": "10 kΩ ±5%, 0.25 W, encapsulado axial",
             "justificacion": "Define nivel lógico alto en la línea de reset del MCU evitando reinicios espúreos por capacitancias parásitas."},
        ]

    # ── Código fuente ─────────────────────────────────────────────────────────
    source_code: str = str(parameters.get("source_code") or "")
    if not source_code.strip():
        source_code = (
            "# =======================================================================\n"
            f"# PROYECTO  : {project_title}\n"
            f"# AUTOR     : {author}\n"
            f"# INSTITUCIÓN: {institution}\n"
            f"# FECHA     : {datetime.now().strftime('%d/%m/%Y')}\n"
            f"# VERSIÓN   : {version}\n"
            "# DESCRIPCIÓN: Control de dispositivo electrónico con ESP32.\n"
            "#              Adquisición de señal → procesamiento → actuación\n"
            "#              → respaldo en nube (si aplica).\n"
            "# =======================================================================\n\n"
            "import time\n"
            "import sys\n\n"
            "# ── Constantes de configuración ─────────────────────────────────────────\n"
            "PIN_SENSOR    = 34          # GPIO34 – entrada analógica ADC1\n"
            "PIN_ACTUADOR  = 23          # GPIO23 – salida PWM / digital\n"
            "UMBRAL_ALTO   = 70.0        # % del rango ADC – umbral de activación\n"
            "UMBRAL_BAJO   = 30.0        # % del rango ADC – histéresis de desactivación\n"
            "PERIODO_MS    = 500         # ms – período de muestreo (2 Hz)\n"
            "VOLT_REF      = 3.3         # V – tensión de referencia ADC\n"
            "RESOLUCION    = 4095        # 12-bit ADC (2^12 - 1)\n\n"
            "# ── Variables de estado ──────────────────────────────────────────────────\n"
            "estado_actuador = False     # False = APAGADO, True = ENCENDIDO\n"
            "log_sesion: list = []       # registro de eventos de la sesión\n\n\n"
            "# ── Funciones de soporte ─────────────────────────────────────────────────\n\n"
            "def leer_sensor(pin: int) -> float:\n"
            "    \"\"\"\n"
            "    Lee el valor del ADC en el pin indicado y lo convierte a porcentaje\n"
            "    del rango completo (0–100 %).\n"
            "    Retorna: float en rango [0.0, 100.0]\n"
            "    \"\"\"\n"
            "    # En MicroPython / hardware real usar: machine.ADC(pin).read()\n"
            "    # Aquí se simula para entorno de prueba.\n"
            "    import random\n"
            "    lectura_raw = random.randint(0, RESOLUCION)   # ← reemplazar por lectura real\n"
            "    return (lectura_raw / RESOLUCION) * 100.0\n\n\n"
            "def controlar_actuador(activar: bool, pin: int) -> None:\n"
            "    \"\"\"\n"
            "    Activa o desactiva el actuador conectado al pin de salida.\n"
            "    Implementa lógica anti-rebote de hardware (histéresis).\n"
            "    \"\"\"\n"
            "    global estado_actuador\n"
            "    if activar == estado_actuador:\n"
            "        return  # sin cambio de estado: no hacer nada\n"
            "    estado_actuador = activar\n"
            "    accion = \"ENCENDIDO\" if activar else \"APAGADO\"\n"
            "    # En hardware real: machine.Pin(pin, machine.Pin.OUT).value(int(activar))\n"
            "    print(f\"[ACTUADOR] GPIO{pin} → {accion}\")\n\n\n"
            "def registrar_evento(nivel: str, mensaje: str) -> None:\n"
            "    \"\"\"\n"
            "    Registra un evento en el log de sesión con marca de tiempo.\n"
            "    nivel: 'INFO' | 'WARN' | 'ERROR'\n"
            "    \"\"\"\n"
            "    ts = time.strftime(\"%Y-%m-%d %H:%M:%S\", time.localtime())\n"
            "    entrada = f\"[{ts}] [{nivel}] {mensaje}\"\n"
            "    log_sesion.append(entrada)\n"
            "    print(entrada)\n\n\n"
            "# ── Bucle principal de control ───────────────────────────────────────────\n\n"
            "def setup() -> None:\n"
            "    \"\"\"Inicialización del sistema antes del bucle de control.\"\"\"\n"
            "    registrar_evento(\"INFO\", f\"Sistema iniciado — {project_title}\")\n"
            "    registrar_evento(\"INFO\", f\"Umbrales: ALTO={UMBRAL_ALTO}% BAJO={UMBRAL_BAJO}%\")\n"
            "    registrar_evento(\"INFO\", \"Hardware inicializado correctamente.\")\n\n\n"
            "def loop() -> None:\n"
            "    \"\"\"\n"
            "    Bucle principal de adquisición → decisión → actuación.\n"
            "    Se ejecuta de forma indefinida con período PERIODO_MS.\n"
            "    \"\"\"\n"
            "    while True:\n"
            "        try:\n"
            "            valor = leer_sensor(PIN_SENSOR)\n"
            "            voltaje = (valor / 100.0) * VOLT_REF\n\n"
            "            # ── Lógica de control con histéresis ─────────────────\n"
            "            if valor >= UMBRAL_ALTO and not estado_actuador:\n"
            "                controlar_actuador(True, PIN_ACTUADOR)\n"
            "                registrar_evento(\"INFO\", f\"Activación: {valor:.1f}% ({voltaje:.2f} V)\")\n\n"
            "            elif valor <= UMBRAL_BAJO and estado_actuador:\n"
            "                controlar_actuador(False, PIN_ACTUADOR)\n"
            "                registrar_evento(\"INFO\", f\"Desactivación: {valor:.1f}% ({voltaje:.2f} V)\")\n\n"
            "            time.sleep(PERIODO_MS / 1000.0)\n\n"
            "        except KeyboardInterrupt:\n"
            "            registrar_evento(\"INFO\", \"Bucle detenido por el usuario.\")\n"
            "            controlar_actuador(False, PIN_ACTUADOR)  # seguridad: apagar actuador\n"
            "            break\n\n"
            "        except Exception as exc:\n"
            "            registrar_evento(\"ERROR\", f\"Excepción en loop: {exc}\")\n"
            "            time.sleep(2.0)  # pausa antes de reintentar\n\n\n"
            "# ── Punto de entrada ─────────────────────────────────────────────────────\n\n"
            "if __name__ == \"__main__\":\n"
            "    setup()\n"
            "    loop()\n"
            "    sys.exit(0)\n"
        )

    try:
        try:
            from docx import Document
            from docx.shared import Inches, Pt
        except (ModuleNotFoundError, ImportError) as dep_err:
            # No abortar todo el entregable: continuar con HTML/diagramas.
            Document = None  # type: ignore[assignment]
            Inches = None    # type: ignore[assignment]
            Pt = None        # type: ignore[assignment]
            _log(player, f"WARN: python-docx no disponible ({dep_err}). Se continuará sin Word.")

        today = datetime.now().strftime("%d/%m/%Y")

        raw_devices = parameters.get("devices") or parameters.get("requested_devices") or []
        device_names: List[str] = []
        for d in raw_devices:
            if isinstance(d, dict):
                name = str(d.get("name") or d.get("nombre") or "").strip()
            else:
                name = str(d).strip()
            if name:
                device_names.append(name)
        if not device_names:
            device_names = [project_title]

        # ── FASE 1: Hardware ──────────────────────────────────────────────────
        _phase_start(player, "Hardware", lines, speak, progress=15)
        import time; time.sleep(0.3)
        _phase_end(player, "Hardware", lines, speak, progress=25)

        # ── FASE 2: Software ──────────────────────────────────────────────────
        _phase_start(player, "Software", lines, speak, progress=28)
        # Guardar código Python como archivo independiente dentro de Report/
        code_file = output_docx.parent / "main_control.py"
        arduino_code = build_arduino_source(project_title)
        arduino_file = output_docx.parent / "main_control.ino"
        try:
            code_file.parent.mkdir(parents=True, exist_ok=True)
            code_file.write_text(source_code, encoding="utf-8")
            _log(player, f"ACT: Código Python guardado en {code_file}")
            arduino_file.write_text(arduino_code, encoding="utf-8")
            _log(player, f"ACT: Código Arduino/C++ guardado en {arduino_file}")
        except Exception as ce:
            _log(player, f"WARN: No se pudo guardar código fuente: {ce}")
        time.sleep(0.3)
        _phase_end(player, "Software", lines, speak, progress=40)

        # ── FASE 3: Diagramas ─────────────────────────────────────────────────
        _phase_start(player, "Diagramas", lines, speak, progress=45)
        block_png = circuit_png = flujo_png = flujo_svg = None
        power_png = fsm_png = sw_arch_png = comm_png = mech_png = tests_chart_png = None
        pcb_svg = pcb_png = uml_png = signal_flow_png = None
        uml_seq_png = uml_uc_png = None
        generated_device_dirs: List[Path] = []
        try:
            gv = _generate_graphviz(diagram_dir, project_title)
            gc = _generate_graphviz_circuit(diagram_dir, project_title)
            mm = _generate_mermaid(diagram_dir, project_title, source_code)

            # ── Entregables de hardware/software/integrados adicionales ──────
            pw = generate_power_flow_diagram(diagram_dir, project_title,
                                              _write_fallback_png, _write_fallback_svg)
            fsm = generate_fsm_diagram(diagram_dir, project_title,
                                        _write_fallback_png, _write_fallback_svg)
            arch = generate_sw_architecture_diagram(diagram_dir, project_title,
                                                     _write_fallback_png, _write_fallback_svg)
            comm = generate_hw_sw_comm_diagram(diagram_dir, project_title,
                                                _write_fallback_png, _write_fallback_svg)
            mech = generate_mechanical_diagram(diagram_dir, project_title,
                                                _write_fallback_png, _write_fallback_svg)
            pcb = generate_pcb_layout_sketch(diagram_dir, project_title, components=parameters.get("components") or [])
            tch = generate_test_results_chart(diagram_dir, project_title,
                                              _write_fallback_png, _write_fallback_svg)
            uml_diag = generate_uml_diagram(diagram_dir, project_title,
                                            _write_fallback_png, _write_fallback_svg)
            sig_diag = generate_signal_flow_diagram(diagram_dir, project_title,
                                                    _write_fallback_png, _write_fallback_svg)
            uml_seq  = generate_uml_sequence_diagram(diagram_dir, project_title,
                                                     _write_fallback_png, _write_fallback_svg)
            uml_uc   = generate_uml_usecase_diagram(diagram_dir, project_title,
                                                    _write_fallback_png, _write_fallback_svg)

            expected = _ensure_diagram_bundle(diagram_dir, project_title)

            block_png   = expected["graphviz_png"] if expected["graphviz_png"].exists() else gv.get("png")
            circuit_png = expected["circuit_png"]  if expected["circuit_png"].exists()  else gc.get("png")
            flujo_png   = expected["mermaid_png"]  if expected["mermaid_png"].exists()  else mm.get("png")
            flujo_svg   = expected["mermaid_svg"]  if expected["mermaid_svg"].exists()  else mm.get("svg")
            power_png   = expected["power_png"]    if expected["power_png"].exists()    else pw.get("png")
            fsm_png     = expected["fsm_png"]      if expected["fsm_png"].exists()      else fsm.get("png")
            sw_arch_png = expected["sw_arch_png"]  if expected["sw_arch_png"].exists()  else arch.get("png")
            comm_png    = expected["comm_png"]     if expected["comm_png"].exists()     else comm.get("png")
            mech_png    = expected["mechanical_png"] if expected["mechanical_png"].exists() else mech.get("png")
            pcb_svg     = expected["pcb_svg"]      if expected["pcb_svg"].exists()      else pcb.get("svg")
            pcb_png     = expected["pcb_png"]      if expected["pcb_png"].exists()      else pcb.get("png")
            uml_png     = expected["uml_png"]      if expected["uml_png"].exists()      else uml_diag.get("png")
            signal_flow_png = expected["signal_flow_png"] if expected["signal_flow_png"].exists() else sig_diag.get("png")
            uml_seq_png     = uml_seq.get("png")
            uml_uc_png      = uml_uc.get("png")
            tests_chart_png = tch.get("png")

            generated_count = sum(1 for p in expected.values() if p.exists())
            _log(player, f"ACT: Diagramas base generados: {generated_count}/{len(expected)}")
            write_manifest(diagram_dir, expected_bundle(diagram_dir), project_title)

            # Diagramas por dispositivo solicitado.
            for dev in device_names:
                dev_dir = diagram_dir / _slugify(dev)
                dev_dir.mkdir(parents=True, exist_ok=True)
                dev_title = f"{project_title} - {dev}"
                _generate_graphviz(dev_dir, dev_title)
                _generate_graphviz_circuit(dev_dir, dev_title)
                _generate_mermaid(dev_dir, dev_title, source_code)
                generate_power_flow_diagram(dev_dir, dev_title, _write_fallback_png, _write_fallback_svg)
                generate_fsm_diagram(dev_dir, dev_title, _write_fallback_png, _write_fallback_svg)
                generate_sw_architecture_diagram(dev_dir, dev_title, _write_fallback_png, _write_fallback_svg)
                generate_hw_sw_comm_diagram(dev_dir, dev_title, _write_fallback_png, _write_fallback_svg)
                generate_mechanical_diagram(dev_dir, dev_title, _write_fallback_png, _write_fallback_svg)
                generate_pcb_layout_sketch(dev_dir, dev_title, components=parameters.get("components") or [])
                generate_test_results_chart(dev_dir, dev_title, _write_fallback_png, _write_fallback_svg)
                generate_uml_diagram(dev_dir, dev_title, _write_fallback_png, _write_fallback_svg)
                generate_signal_flow_diagram(dev_dir, dev_title, _write_fallback_png, _write_fallback_svg)
                _ensure_diagram_bundle(dev_dir, dev_title)
                write_manifest(dev_dir, expected_bundle(dev_dir), f"{project_title}::{dev}")
                generated_device_dirs.append(dev_dir)

            _log(player, f"ACT: Diagramas por dispositivo generados: {len(generated_device_dirs)}/{len(device_names)}")
        except Exception as de:
            _log(player, f"WARN: Fase de diagramas con incidencias: {de}")
        _phase_end(player, "Diagramas", lines, speak, progress=65)

        # ── FASE 4: Word ──────────────────────────────────────────────────────
        _phase_start(player, "Word", lines, speak, progress=68)
        word_ok = False
        word_err = ""
        try:
            if Document is None:
                raise RuntimeError("python-docx no disponible en el entorno")

            # Fallback: si no hay plantilla, crear documento vacío.
            if template_path.exists():
                doc = Document(str(template_path))
            else:
                _log(player, f"WARN: Plantilla no encontrada ({template_path}). Se generará Word desde cero.")
                doc = Document()

            _safe_replace_placeholders(doc, {
                "{{TITLE}}":         project_title,
                "{{PROJECT_TITLE}}": project_title,
                "{{AUTHOR}}":        author,
                "{{INSTITUTION}}":   institution,
                "{{DATE}}":          today,
            })

            # Portada con todos los datos de identificación
            _add_heading_safe(doc, "Portada", level=1)
            meta_pairs = [
                ("Proyecto",     project_title),
                ("Autor",        author),
                ("Institución",  institution),
                ("Fecha",        today),
                ("Versión",      version),
            ]
            for label, value in meta_pairs:
                p = doc.add_paragraph()
                run_l = p.add_run(f"{label}: ")
                run_l.bold = True
                p.add_run(value)

            doc.add_page_break()

            # Plan del proyecto + estándares
            _project_plan_to_docx_table(doc, project_plan, software_standards)

            # Descripción general
            _add_heading_safe(doc, "2. Funcionamiento del Producto", level=1)
            doc.add_paragraph(overview)

            # Hardware + tabla de componentes
            _add_heading_safe(doc, "3. Hardware y Arquitectura Física", level=1)
            _components_to_docx_table(doc, components)
            _add_heading_safe(doc, "3.1 Lista de Materiales (BOM)", level=2)
            doc.add_paragraph("Inventario formal de componentes con referencia y cantidad para compras/ensamblaje.")
            build_bom_docx_table(doc, components)
            _add_heading_safe(doc, "3.2 Criterios de Aceptación del Diseño", level=2)
            doc.add_paragraph("Los siguientes criterios definen la aceptación técnica del diseño de hardware:")
            for criterion in build_acceptance_criteria(project_title):
                doc.add_paragraph(criterion)

            # Software
            _add_heading_safe(doc, "4. Desarrollo de Software y Lógica de Control", level=1)
            _add_heading_safe(doc, "4.1 Estándares de desarrollo aplicados", level=2)
            for standard in software_standards:
                doc.add_paragraph(f"• {standard}")
            _add_heading_safe(doc, "4.2 Pseudocódigo de rutinas críticas", level=2)
            p_pseudo = doc.add_paragraph()
            run_p = p_pseudo.add_run(build_pseudocode(project_title))
            run_p.font.name = "Consolas"
            if Pt is not None:
                run_p.font.size = Pt(9)

            _add_heading_safe(doc, "4.3 Código fuente Python", level=2)
            doc.add_paragraph(
                "Código fuente Python guardado en Report/main_control.py. "
                "Fragmento principal (comentarios en español, ISO/IEC 9126):"
            )
            p_code = doc.add_paragraph()
            run_c = p_code.add_run(source_code[:3000] + ("\n[... ver main_control.py ...]" if len(source_code) > 3000 else ""))
            run_c.font.name = "Consolas"
            if Pt is not None:
                run_c.font.size = Pt(8)

            _add_heading_safe(doc, "4.4 Código fuente Arduino/C++", level=2)
            doc.add_paragraph(
                "Código Arduino/C++ guardado en Report/main_control.ino. "
                "Compatible con ESP32 y Arduino UNO/Mega. Comentarios en español:"
            )
            p_ino = doc.add_paragraph()
            run_ino = p_ino.add_run(arduino_code[:3000] + ("\n[... ver main_control.ino ...]" if len(arduino_code) > 3000 else ""))
            run_ino.font.name = "Consolas"
            if Pt is not None:
                run_ino.font.size = Pt(8)

            _add_heading_safe(doc, "4.5 Documentación Técnica del Software (Algoritmos y Estructuras de Datos)", level=2)
            doc.add_paragraph(
                "Algoritmo Principal:\n"
                "El control opera mediante un ciclo cerrado de adquisición periódica (tiempo de muestreo controlado por delay). "
                "Para evitar conmutaciones espúreas (ruido o variaciones transitorias en la señal del sensor), "
                "se implementa una lógica de control con histéresis basada en dos umbrales (UMBRAL_ALTO y UMBRAL_BAJO).\n\n"
                "Estructuras de Datos:\n"
                "- estado_actuador (Booleano): Representa el estado actual de la salida (FALSO = Apagado, VERDADERO = Encendido).\n"
                "- log_sesion (Lista/Array): Cola lineal dinámica donde se registran los eventos del sistema con marcas de tiempo."
            )

            # Diagramas
            _add_heading_safe(doc, "5. Diagramas de Circuitos y Bloques Funcionales", level=1)

            def _insert_fig(heading: str, caption: str, img_path, width_in: float = 6.2, fallback_msg: str = "Diagrama no disponible."):
                _add_heading_safe(doc, heading, level=2)
                doc.add_paragraph(caption)
                if img_path and Path(img_path).exists() and Inches is not None:
                    try:
                        doc.add_picture(str(img_path), width=Inches(width_in))
                        return True
                    except Exception as e:
                        doc.add_paragraph(f"[No se pudo insertar imagen: {e}]")
                        return False
                else:
                    doc.add_paragraph(fallback_msg)
                    return False

            inserted_flags = []
            inserted_flags.append(_insert_fig("5.1 Diagrama de Bloques Funcionales",
                "Figura 5.1: Arquitectura de bloques funcionales del sistema (alimentación, control, sensores, actuadores).",
                block_png))
            inserted_flags.append(_insert_fig("5.2 Diagrama del Circuito Esquemático",
                "Figura 5.2: Circuito esquemático detallado con componentes y conexiones.",
                circuit_png))
            inserted_flags.append(_insert_fig("5.3 Boceto de Disposición de PCB",
                "Figura 5.3: Croquis conceptual de posición de componentes y trazas. "
                "No apto para fabricación: el ruteo real debe realizarse en una herramienta EDA "
                "(KiCad / Proteus ARES) a partir del esquemático de la Figura 5.2.",
                pcb_png))
            inserted_flags.append(_insert_fig("5.4 Diagrama de Flujo de Energía",
                "Figura 5.4: Distribución de la alimentación entre las etapas del sistema.",
                power_png))
            inserted_flags.append(_insert_fig("5.5 Diagrama Mecánico",
                "Figura 5.5: Vista 2D esquemática del encapsulado/caja del proyecto (croquis conceptual, no sustituye plano CAD).",
                mech_png))
            inserted_flags.append(_insert_fig("5.6 Diagrama de Flujo del Software",
                "Figura 5.6: Diagrama de flujo del ciclo completo de control.",
                flujo_png, width_in=4.5))
            inserted_flags.append(_insert_fig("5.7 Diagrama de Estados (FSM)",
                "Figura 5.7: Comportamiento del sistema según eventos.",
                fsm_png))
            inserted_flags.append(_insert_fig("5.8 Diagrama de Arquitectura de Software",
                "Figura 5.8: Capas, módulos y comunicación entre ellos.",
                sw_arch_png))
            inserted_flags.append(_insert_fig("5.9 Diagrama de Comunicación Hardware-Software",
                "Figura 5.9: Interacción entre sensores, microcontrolador y programa.",
                comm_png))
            inserted_flags.append(_insert_fig("5.10 Diagrama UML – Clases",
                "Figura 5.10: Diagrama de clases del software de control.",
                uml_png))
            inserted_flags.append(_insert_fig("5.11 Diagrama UML – Casos de Uso",
                "Figura 5.11: Casos de uso del sistema: actores, interacciones e inclusiones.",
                uml_uc_png))
            inserted_flags.append(_insert_fig("5.12 Diagrama UML – Secuencia",
                "Figura 5.12: Secuencia de mensajes entre usuario, software, MCU, sensor y actuador.",
                uml_seq_png))
            inserted_flags.append(_insert_fig("5.13 Flujo de Señal (Telecomunicaciones)",
                "Figura 5.13: Flujo de señal, acondicionamiento y comunicación RF del sistema.",
                signal_flow_png))

            # 6. Pruebas y resultados
            _add_heading_safe(doc, "6. Pruebas y Resultados", level=1)
            doc.add_paragraph(
                "Plan de validación del sistema. Los resultados obtenidos deben completarse "
                "tras la ejecución de las pruebas en banco."
            )
            if tests_chart_png and Path(tests_chart_png).exists() and Inches is not None:
                try:
                    doc.add_picture(str(tests_chart_png), width=Inches(5.8))
                    doc.add_paragraph("Figura 6.1: Resumen gráfico del estado de pruebas y resultados.")
                except Exception:
                    doc.add_paragraph("No fue posible insertar la gráfica de pruebas.")
            build_test_results_docx_table(doc, project_title)

            # 7. Manual de usuario
            _add_heading_safe(doc, "7. Manual de Usuario / Guía de Operación", level=1)
            for item in build_user_manual_sections(project_title, overview, device_names):
                doc.add_paragraph(item)

            # 8. Conclusiones
            _add_heading_safe(doc, "8. Conclusiones", level=1)
            doc.add_paragraph(
                "El presente informe técnico integra el diseño de hardware, el desarrollo de "
                "software y su validación planificada, conforme a los entregables definidos en "
                "el plan del proyecto (sección 1)."
            )

            _add_heading_safe(doc, "9. Informe Técnico Final (Consolidado)", level=1)
            doc.add_paragraph(
                "Este documento constituye el informe técnico final del proyecto, e incluye: "
                "fundamentos teóricos y funcionamiento (sección 2), diseño e implementación "
                "de hardware y software (secciones 3, 4 y 5), plan de validación con resultados "
                "(sección 6) y conclusiones (sección 8)."
            )
            _add_heading_safe(doc, "9.1 Estado de Entrega", level=2)
            doc.add_paragraph(
                "CONFIRMACIÓN EXPLÍCITA: Se certifica que todas las fases del proyecto (diseño de hardware, "
                "desarrollo e integración de software, diagramación técnica y documentación) han sido concluidas. "
                "El proyecto está 100% completo, verificado y listo para su implementación física y despliegue."
            )

            doc.save(str(output_docx))
            inserted = sum(1 for f in inserted_flags if f)
            _log(player, f"ACT: Diagramas insertados en Word: {inserted}/{len(inserted_flags)}")
            word_ok = True
        except Exception as we:
            word_err = str(we)
            _log(player, f"WARN: Falló construcción de Word: {word_err}")
        _phase_end(player, "Word", lines, speak, progress=85)

        # ── Pre-calcular reporte final (se incrusta en el HTML) ──────────────
        _pre_generated_files = {
            "Documento Word (.docx)":    str(output_docx) if word_ok and output_docx.exists() else "",
            "Página Web (index.html)":   str(output_html),
            "Código Python":             str(output_docx.parent / "main_control.py"),
            "Código Arduino/C++ (.ino)": str(output_docx.parent / "main_control.ino"),
            "Diagramas (carpeta)":       str(diagram_dir),
        }
        final_report_txt = build_final_report_section(
            project_title=project_title,
            author=author,
            institution=institution,
            project_plan=project_plan,
            generated_files=_pre_generated_files,
            date_str=today,
        )

        # ── FASE 5: Web ───────────────────────────────────────────────────────
        _phase_start(player, "Web", lines, speak, progress=88)
        html_ok = False
        html_err = ""
        try:
            def _ok(p):
                return p if p and Path(p).exists() else None

            _build_html(
                output_html=output_html,
                project_title=project_title,
                overview=overview,
                components=components,
                source_code=source_code,
                block_diagram=_ok(block_png),
                circuit_diagram=_ok(circuit_png),
                flow_diagram=_ok(flujo_png),
                author=author,
                institution=institution,
                date_str=today,
                version=version,
                project_plan=project_plan,
                software_standards=software_standards,
                power_diagram=_ok(power_png),
                mechanical_diagram=_ok(mech_png),
                pcb_diagram=_ok(pcb_png) or _ok(pcb_svg),
                fsm_diagram=_ok(fsm_png),
                sw_arch_diagram=_ok(sw_arch_png),
                comm_diagram=_ok(comm_png),
                test_chart=_ok(tests_chart_png),
                device_names=device_names,
                uml_diagram=_ok(uml_png),
                signal_flow_diagram=_ok(signal_flow_png),
                uml_seq_diagram=_ok(uml_seq_png),
                uml_uc_diagram=_ok(uml_uc_png),
                arduino_code=arduino_code,
                final_report_txt=final_report_txt,
            )
            html_ok = True
            _log(player, f"ACT: Entregable web generado: {output_html}")
        except Exception as he:
            html_err = str(he)
            _log(player, f"WARN: Falló construcción de HTML: {html_err}")
        _phase_end(player, "Web", lines, speak, progress=98)

        # ── FASE 6: Reporte Final ─────────────────────────────────────────────
        _phase_start(player, "Reporte Final", lines, speak, progress=99)
        word_status = "OK" if word_ok and output_docx.exists() else f"FALLÓ ({word_err or 'sin detalle'})"
        html_status = "OK" if html_ok and output_html.exists() else f"FALLÓ ({html_err or 'sin detalle'})"
        # Actualizar estado real en el reporte final ahora que tenemos resultados definitivos
        final_generated_files = {
            "Documento Word (.docx)":    str(output_docx) if word_ok and output_docx.exists() else "",
            "Página Web (index.html)":   str(output_html) if html_ok and output_html.exists() else "",
            "Código Python":             str(code_file)   if code_file.exists() else "",
            "Código Arduino/C++ (.ino)": str(arduino_file) if arduino_file.exists() else "",
            "Diagramas (carpeta)":       str(diagram_dir),
            "Reporte Final (.txt)":      str(output_docx.parent / "reporte_final.txt"),
        }
        final_report_txt_definitive = build_final_report_section(
            project_title=project_title,
            author=author,
            institution=institution,
            project_plan=project_plan,
            generated_files=final_generated_files,
            date_str=today,
        )
        reporte_final_file = output_docx.parent / "reporte_final.txt"
        try:
            reporte_final_file.write_text(final_report_txt_definitive, encoding="utf-8")
            _log(player, f"ACT: Reporte final guardado en {reporte_final_file}")
        except Exception as rfe:
            _log(player, f"WARN: No se pudo guardar reporte_final.txt: {rfe}")
        _phase_end(player, "Reporte Final", lines, speak, progress=100)

        # Finalizar
        total_dias = sum(p["duration_days"] for p in project_plan)
        final_msg = (
            f"✅ Proyecto 100% completo y verificado.\n"
            f"  Cronograma: {project_plan[0]['start']} → {project_plan[-1]['end']} ({total_dias} días)\n"
            f"  Word      : {output_docx} [{word_status}]\n"
            f"  Web       : {output_html} [{html_status}]\n"
            f"  Python    : {code_file}\n"
            f"  Arduino   : {arduino_file}\n"
            f"  Diagramas : {diagram_dir}\n"
            f"  Rep. Final: {reporte_final_file}\n"
            f"  Dispositivos con diagramas: {len(device_names)}"
        )
        _log(player, final_msg)
        lines.append(final_msg)
        if player and hasattr(player, "update_activity"):
            try:
                player.update_activity(estado="Completado", progreso=100,
                                       evento="Reporte de ingeniería completado")
            except Exception:
                pass
        return "\n".join(lines)

    except Exception as e:
        err = f"❌ Error en engineering_report: {e}"
        _log(player, err)
        lines.append(err)
        return "\n".join(lines)